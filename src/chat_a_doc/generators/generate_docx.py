"""DOCX generator from markdown content.

Uses python-docx to create DOCX documents from markdown, with
support for template documents (reference_doc).
"""

import re
from io import BytesIO
from pathlib import Path

import markdown
from docx import Document
from docx.shared import Pt


def generate_docx(markdown_content: str, reference_doc: str, **options) -> bytes:
    """Generate DOCX from markdown content.

    Args:
        markdown_content: Markdown content as string
        reference_doc: Required path to template DOCX file
        **options: Additional options
            - title: Document title (for metadata)

    Returns:
        DOCX bytes

    Raises:
        ValueError: If DOCX generation fails
        FileNotFoundError: If reference_doc file doesn't exist

    """
    try:
        # Step 1: Load template document
        template_path = Path(reference_doc)
        if not template_path.exists():
            raise FileNotFoundError(f"Template file not found: {reference_doc}")

        # Open template document
        # Note: Template should be pre-cleared by user (no placeholder text)
        # The generator uses the template's styles but does not clear content
        doc = Document(str(template_path))

        # Step 2: Parse markdown to extract structure
        # Step 3: Parse markdown AST directly for better structure
        # Use markdown's tree processor to get structured content
        extensions = ["tables", "fenced_code", "nl2br"]
        md = markdown.Markdown(extensions=extensions)
        md.convert(markdown_content)

        # Step 4: Add content to document
        # For now, use a simpler approach: parse markdown line by line
        # and convert to DOCX elements
        add_markdown_to_docx(doc, markdown_content)

        # Step 5: Save to bytes
        output = BytesIO()
        doc.save(output)

        return output.getvalue()

    except Exception as e:
        raise ValueError(f"Failed to generate DOCX from markdown: {e}") from e


def add_markdown_to_docx(doc: Document, markdown_content: str) -> None:
    """Add markdown content to a DOCX document.

    Parses markdown and adds appropriate DOCX elements.

    Args:
        doc: python-docx Document object
        markdown_content: Markdown content to add

    """
    lines = markdown_content.split("\n")
    in_code_block = False
    code_block_lines = []

    for line in lines:
        # Handle code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                if code_block_lines:
                    code_para = doc.add_paragraph("".join(code_block_lines))
                    code_para.style = "No Spacing"
                    # Apply monospace font
                    for run in code_para.runs:
                        run.font.name = "Courier New"
                        run.font.size = Pt(10)
                code_block_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            continue

        if in_code_block:
            code_block_lines.append(line + "\n")
            continue

        # Handle headers
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("##### "):
            doc.add_heading(line[6:].strip(), level=5)
        elif line.startswith("###### "):
            doc.add_heading(line[7:].strip(), level=6)
        # Handle horizontal rules
        elif line.strip() == "---" or line.strip() == "***":
            doc.add_paragraph().add_run().add_break()
        # Handle lists
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            item_text = line.strip()[2:].strip()
            # Remove markdown formatting from list items
            item_text = remove_markdown_formatting(item_text)
            doc.add_paragraph(item_text, style="List Bullet")
        elif line.strip().startswith(("1. ", "2. ", "3. ", "4. ", "5. ")):
            # Numbered list (simplified - only handles single digits)
            item_text = line.strip().split(". ", 1)[1] if ". " in line else line.strip()
            item_text = remove_markdown_formatting(item_text)
            doc.add_paragraph(item_text, style="List Number")
        # Handle tables (basic markdown table detection)
        elif "|" in line and line.count("|") >= 2:
            # This is a table row (simplified - would need more parsing for full table support)
            # For now, skip table separator rows
            if not re.match(r"^[\s\|:\-]+$", line):
                cells = [cell.strip() for cell in line.split("|") if cell.strip()]
                if cells:
                    # Add as simple paragraph for now
                    # Full table support would require more complex parsing
                    doc.add_paragraph(" | ".join(cells))
        # Handle empty lines
        elif not line.strip():
            doc.add_paragraph()
        # Handle regular paragraphs
        else:
            # Check if line has formatting that needs to be preserved
            has_formatting = bool(re.search(r"\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_|`.*?`|\[.*?\]\(.*?\)", line))
            if has_formatting:
                # Create empty paragraph and apply formatting (preserves original markdown)
                para = doc.add_paragraph()
                apply_formatting_to_paragraph(para, line)
            else:
                # No formatting, just add plain text
                para_text = remove_markdown_formatting(line)
                if para_text.strip():
                    doc.add_paragraph(para_text)


def remove_markdown_formatting(text: str) -> str:
    """Remove markdown formatting from text.

    Args:
        text: Text with markdown formatting

    Returns:
        Plain text

    """
    # Remove bold/italic
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"_(.*?)_", r"\1", text)
    # Remove links
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    # Remove inline code
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def apply_formatting_to_paragraph(para, original_line: str) -> None:
    """Apply formatting (bold, italic) to a paragraph based on markdown.

    Args:
        para: python-docx Paragraph object (should be empty)
        original_line: Original markdown line with formatting

    """
    # Paragraph should already be empty when this function is called
    # This ensures we don't lose any text that was added

    # Split by markdown formatting and add runs
    # This is simplified - a full implementation would parse more carefully
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_|`.*?`)", original_line)

    for part in parts:
        if not part:
            continue

        # Check formatting BEFORE removing markers (so patterns can match)
        is_bold = bool(re.match(r"\*\*.*?\*\*", part) or re.match(r"__.*?__", part))
        is_italic = bool(re.match(r"\*.*?\*", part) or re.match(r"_.*?_", part))
        is_code = bool(re.match(r"`.*?`", part))

        # Remove formatting markers and add text to run
        run = para.add_run(remove_markdown_formatting(part))

        # Apply formatting based on original part (before stripping)
        if is_bold:
            run.bold = True
        elif is_italic:
            run.italic = True
        elif is_code:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
